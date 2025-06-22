from docx import Document

def generate_docx(info_dict, output_path):
    doc = Document()
    doc.add_heading("사양 추출 결과", level=1)
    for key, value in info_dict.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(output_path)